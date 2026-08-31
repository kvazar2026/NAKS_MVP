"""One-off generator for ``backend/templates/demo_equipment_application.docx``
— the single ``(demo, equipment)`` ``template registry`` entry (spec).

Kept as source (not just the committed binary ``.docx``) so the template's
placeholders and copy stay reviewable as a text diff. Re-run whenever the
template's content needs to change; the placeholder names must match the
keys ``app/services/document_generation.py``'s ``build_template_context``
produces.

Usage (from the ``backend`` directory, with the venv active):
    python scripts/generate_demo_template.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_PATH = Path(__file__).resolve().parents[1] / "templates" / "demo_equipment_application.docx"


def build() -> Document:
    doc = Document()

    title = doc.add_heading("Заявка на аттестацию сварочного оборудования", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    disclaimer = doc.add_paragraph()
    disclaimer_run = disclaimer.add_run("Демонстрационный макет, не является официальным бланком АЦ.")
    disclaimer_run.bold = True
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Аттестационный центр: {{ ac_name }}")
    doc.add_paragraph("Группа ОПО: {{ opo_group }}")
    doc.add_paragraph("Регион проведения аттестации: {{ region }}")

    doc.add_heading("1. Сведения об организации", level=2)
    doc.add_paragraph("Наименование организации: {{ company_name }}")
    doc.add_paragraph("ИНН: {{ inn }}")
    doc.add_paragraph("Адрес: {{ address }}")

    doc.add_heading("2. Контактное лицо", level=2)
    doc.add_paragraph("ФИО: {{ contact_full_name }}")
    doc.add_paragraph("Должность: {{ contact_position }}")
    doc.add_paragraph("Телефон: {{ contact_phone }}")
    doc.add_paragraph("Email: {{ contact_email }}")

    doc.add_heading("3. Сведения об оборудовании", level=2)
    doc.add_paragraph("Тип оборудования: {{ equipment_type }}")
    doc.add_paragraph("Марка: {{ brand }}")
    doc.add_paragraph("Модель: {{ model }}")
    doc.add_paragraph("Изготовитель: {{ manufacturer }}")
    doc.add_paragraph("Способ сварки: {{ welding_method }}")
    doc.add_paragraph("Количество, шт.: {{ quantity }}")
    doc.add_paragraph("Заводские номера: {{ serial_numbers }}")
    doc.add_paragraph("Назначение: {{ purpose }}")

    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "Документ сформирован автоматически в демонстрационном режиме и не подлежит подаче в аттестационный центр."
    )
    footer_run.italic = True

    return doc


def main() -> None:
    document = build()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
